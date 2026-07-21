"""
Обробник протоколів засідань - конвертер у таблицю доручень.

Потрібні бібліотеки:
    pip install python-docx

Запуск:
    python protocol_processor.py
"""

import os
import re
import platform
from datetime import datetime

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# =============================================================================
#  Шрифти для крос-платформного вигляду GUI (Windows / Linux / macOS)
# =============================================================================

def get_ui_font():
    """Підбирає шрифт, який точно є в системі, щоб інтерфейс виглядав
    однаково акуратно на Windows, Linux та macOS."""
    system = platform.system()
    if system == "Windows":
        candidates = ["Segoe UI", "Tahoma", "Arial"]
    elif system == "Darwin":
        candidates = ["Helvetica Neue", "Helvetica", "Arial"]
    else:
        candidates = ["Noto Sans", "DejaVu Sans",
                      "Liberation Sans", "Arial", "Ubuntu"]

    try:
        import tkinter.font as tkfont
        available = set(tkfont.families())
    except Exception:
        available = set()

    for c in candidates:
        if c in available:
            return c
    return "TkDefaultFont"


# =============================================================================
#  Робота зі списком виконавців (окремий .txt файл)
# =============================================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(
    __file__)) if "__file__" in globals() else os.getcwd()
EXECUTORS_FILE = os.path.join(SCRIPT_DIR, "executors.txt")

DEFAULT_EXECUTORS = [
    ("Бочін Сергій",        ["Сергію БОЧІНУ",
     "Сергію Бочіну", "Сергій БОЧІН"]),
    ("Гавриленко Лариса",   ["Ларисі ГАВРИЛЕНКО", "Ларисі Гавриленко"]),
    ("Білогуб Вячеслав",    ["Вячеславу БІЛОГУБУ", "Вячеславу Білогубу"]),
    ("Бука Олексій",        ["Олексію БУКА", "Олексію Буці"]),
    ("Коренкова Юлія",      ["Юлії КОРЕНКОВІЙ", "Юлії Коренковій"]),
    ("Леонтьєв Михайло",    ["Михайлу ЛЕОНТЬЄВУ", "Михайлу Леонтьєву"]),
    ("Залуський Андрій",    ["Андрію ЗАЛУСЬКОМУ", "Андрію Залуському"]),
    ("Іщенко Андрій",       ["Андрію ІЩЕНКО", "Андрію Іщенку"]),
    ("Москаленко Ніна",     ["Ніні МОСКАЛЕНКО", "Ніні Москаленко"]),
    ("Олінкевич Олег",      ["Олегу ОЛІНКЕВИЧ", "Олегу Олінкевичу"]),
    ("Всім присутнім",      ["Всім присутнім"]),
]


def load_executors():
    """Завантажує список виконавців з .txt. Якщо файлу нема - створює
    його зі значеннями за замовчуванням."""
    if not os.path.isfile(EXECUTORS_FILE):
        save_executors(DEFAULT_EXECUTORS)
        return [(c, list(v)) for c, v in DEFAULT_EXECUTORS]

    executors = []
    try:
        with open(EXECUTORS_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "|" not in line:
                    continue
                canonical, variants_str = line.split("|", 1)
                variants = [v.strip()
                            for v in variants_str.split(",") if v.strip()]
                if canonical.strip() and variants:
                    executors.append((canonical.strip(), variants))
    except Exception:
        return [(c, list(v)) for c, v in DEFAULT_EXECUTORS]

    return executors if executors else [(c, list(v)) for c, v in DEFAULT_EXECUTORS]


def save_executors(executors):
    """Зберігає список виконавців у .txt файл."""
    with open(EXECUTORS_FILE, "w", encoding="utf-8") as f:
        f.write("# Формат: Прізвище Ім'я | варіант1, варіант2, варіант3\n")
        f.write("# Варіанти - всі форми, як ім'я може зустрічатись у протоколі\n")
        f.write("# (називний відмінок, давальний відмінок тощо)\n")
        for canonical, variants in executors:
            f.write(f"{canonical}|{', '.join(variants)}\n")


def build_executor_patterns(executors):
    """Формує список (regex-фрагмент, канонічне ім'я), довші варіанти
    йдуть першими, щоб уникнути хибних часткових збігів."""
    patterns = []
    for canonical, variants in executors:
        for variant in variants:
            patterns.append((re.escape(variant), canonical))
    patterns.sort(key=lambda p: len(p[0]), reverse=True)
    return patterns


# =============================================================================
#  Вікно управління виконавцями
# =============================================================================

class ExecutorsManager(tk.Toplevel):
    def __init__(self, parent, executors, on_save):
        super().__init__(parent)
        self.title("Управління виконавцями")
        self.geometry("640x680")
        self.executors = [(c, list(v)) for c, v in executors]
        self.on_save = on_save
        self.transient(parent)
        self.grab_set()
        self._build_ui()
        self._refresh_list()

    def _build_ui(self):
        main = ttk.Frame(self, padding=12)
        main.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main, text="Список виконавців", font=(
            "TkDefaultFont", 11, "bold")).pack(anchor="w")

        list_frame = ttk.Frame(main)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(6, 10))

        columns = ("name", "variants")
        self.tree = ttk.Treeview(
            list_frame, columns=columns, show="headings", height=10)
        self.tree.heading("name", text="Ім'я (як у таблиці)")
        self.tree.heading("variants", text="Форми з протоколу (через кому)")
        self.tree.column("name", width=180, anchor="w")
        self.tree.column("variants", width=380, anchor="w")
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scroll = ttk.Scrollbar(
            list_frame, orient="vertical", command=self.tree.yview)
        scroll.pack(side=tk.LEFT, fill=tk.Y)
        self.tree.configure(yscrollcommand=scroll.set)

        ttk.Button(main, text="Видалити вибраного",
                   command=self._remove_selected).pack(anchor="w")

        add_frame = ttk.LabelFrame(
            main, text="Додати нового виконавця", padding=10)
        add_frame.pack(fill=tk.X, pady=(12, 0))

        ttk.Label(add_frame, text="Ім'я (напр. «Бочін Сергій»):").grid(
            row=0, column=0, sticky="w", pady=3)
        self.name_var = tk.StringVar()
        ttk.Entry(add_frame, textvariable=self.name_var, width=45).grid(
            row=0, column=1, pady=3, sticky="we")

        ttk.Label(add_frame, text="Форми в протоколі, через кому\n(напр. «Сергію БОЧІНУ, Сергію Бочіну»):").grid(
            row=1, column=0, sticky="w", pady=3)
        self.variants_var = tk.StringVar()
        ttk.Entry(add_frame, textvariable=self.variants_var, width=45).grid(
            row=1, column=1, pady=3, sticky="we")

        add_frame.columnconfigure(1, weight=1)

        ttk.Button(add_frame, text="+ Додати виконавця", command=self._add_executor).grid(
            row=2, column=0, columnspan=2, pady=(8, 0))

        btns = ttk.Frame(main)
        btns.pack(fill=tk.X, pady=(14, 0))
        ttk.Button(btns, text="Зберегти і закрити",
                   command=self._save_and_close).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Скасувати", command=self.destroy).pack(
            side=tk.RIGHT, padx=4)

    def _refresh_list(self):
        self.tree.delete(*self.tree.get_children())
        for canonical, variants in self.executors:
            self.tree.insert("", "end", values=(
                canonical, ", ".join(variants)))

    def _add_executor(self):
        name = self.name_var.get().strip()
        variants_raw = self.variants_var.get().strip()
        if not name or not variants_raw:
            messagebox.showwarning(
                "Увага", "Заповніть обидва поля - ім'я та хоча б одну форму з протоколу.")
            return
        variants = [v.strip() for v in variants_raw.split(",") if v.strip()]
        self.executors.append((name, variants))
        self._refresh_list()
        self.name_var.set("")
        self.variants_var.set("")

    def _remove_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        index = self.tree.index(sel[0])
        del self.executors[index]
        self._refresh_list()

    def _save_and_close(self):
        save_executors(self.executors)
        self.on_save(self.executors)
        self.destroy()


# =============================================================================
#  Головний клас програми
# =============================================================================

class ProtocolProcessor:
    def __init__(self, root, ui_font):
        self.root = root
        self.ui_font = ui_font
        self.root.title("Обробник протоколів → таблиця доручень")
        self.root.geometry("1100x700")
        self.root.minsize(950, 600)

        self.input_file = None
        self.output_data = []
        self.executors = load_executors()

        self._setup_style()
        self._build_ui()

    # ------------------------------------------------------------------ UI --
    def _setup_style(self):
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except Exception:
            pass

        base = (self.ui_font, 10)
        bold = (self.ui_font, 10, "bold")

        style.configure(".", font=base)
        style.configure("TButton", font=base, padding=6)
        style.configure("Accent.TButton", font=bold, padding=8)
        style.configure("TLabelframe.Label", font=bold)
        style.configure("Header.TLabel", font=(self.ui_font, 16, "bold"))
        style.configure("SubHeader.TLabel", font=(self.ui_font, 10))
        style.configure("Treeview", font=base, rowheight=26)
        style.configure("Treeview.Heading", font=bold)
        style.configure("Card.TFrame", background="#f4f7fa")

    def _build_ui(self):
        # ---- Верхня панель ----
        top = ttk.Frame(self.root, padding=(16, 14, 16, 8))
        top.pack(fill=tk.X)

        ttk.Label(top, text="📋 Обробник протоколів",
                  style="Header.TLabel").pack(anchor="w")
        ttk.Label(top, text="Перетворює протокол засідання у структуровану таблицю доручень",
                  style="SubHeader.TLabel").pack(anchor="w", pady=(2, 0))

        ttk.Separator(self.root).pack(fill=tk.X, padx=16)

        # ---- Основна область: ліва панель дій + права панель перегляду ----
        body = ttk.Frame(self.root, padding=16)
        body.pack(fill=tk.BOTH, expand=True)
        body.columnconfigure(1, weight=1)
        body.rowconfigure(0, weight=1)

        # ЛІВА ПАНЕЛЬ
        left = ttk.Frame(body, width=300)
        left.grid(row=0, column=0, sticky="ns", padx=(0, 16))
        left.grid_propagate(False)

        step1 = ttk.LabelFrame(left, text="1. Файл протоколу", padding=12)
        step1.pack(fill=tk.X, pady=(0, 12))

        self.file_label = ttk.Label(
            step1, text="Файл не вибрано", wraplength=250, foreground="#666")
        self.file_label.pack(fill=tk.X, pady=(0, 8))

        ttk.Button(step1, text="📂 Вибрати файл .docx",
                   command=self.select_file).pack(fill=tk.X)

        step2 = ttk.LabelFrame(left, text="2. Обробка", padding=12)
        step2.pack(fill=tk.X, pady=(0, 12))

        ttk.Button(step2, text="⚙ Обробити файл", style="Accent.TButton",
                   command=self.process_file).pack(fill=tk.X)

        self.status_label = ttk.Label(
            step2, text="", foreground="#2e7d32", wraplength=250)
        self.status_label.pack(fill=tk.X, pady=(8, 0))

        step3 = ttk.LabelFrame(left, text="3. Результат", padding=12)
        step3.pack(fill=tk.X, pady=(0, 12))

        ttk.Button(step3, text="💾 Зберегти таблицю (.docx)",
                   command=self.save_result).pack(fill=tk.X, pady=(0, 6))
        ttk.Button(step3, text="🗑 Очистити все",
                   command=self.clear_all).pack(fill=tk.X)

        step4 = ttk.LabelFrame(left, text="Виконавці", padding=12)
        step4.pack(fill=tk.X)

        ttk.Button(step4, text="👥 Управління виконавцями",
                   command=self.open_executors_manager).pack(fill=tk.X)
        ttk.Label(step4, text=f"Файл: {os.path.basename(EXECUTORS_FILE)}",
                  foreground="#666", font=(self.ui_font, 8)).pack(anchor="w", pady=(6, 0))

        # ПРАВА ПАНЕЛЬ - ПЕРЕГЛЯД РЕЗУЛЬТАТУ
        right = ttk.LabelFrame(
            body, text="Попередній перегляд доручень", padding=10)
        right.grid(row=0, column=1, sticky="nsew")
        right.rowconfigure(0, weight=1)
        right.columnconfigure(0, weight=1)

        columns = ("num", "deadline", "executor", "text")
        self.tree = ttk.Treeview(right, columns=columns, show="headings")
        self.tree.heading("num", text="№")
        self.tree.heading("deadline", text="Термін")
        self.tree.heading("executor", text="Виконавець")
        self.tree.heading("text", text="Доручення")
        self.tree.column("num", width=40, anchor="center", stretch=False)
        self.tree.column("deadline", width=110, anchor="center", stretch=False)
        self.tree.column("executor", width=160, anchor="center", stretch=False)
        self.tree.column("text", width=500, anchor="w")

        self.tree.grid(row=0, column=0, sticky="nsew")

        vscroll = ttk.Scrollbar(right, orient="vertical",
                                command=self.tree.yview)
        vscroll.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=vscroll.set)

        self.tree.bind("<Double-1>", self._show_full_text)

        self.count_label = ttk.Label(
            right, text="Доручень ще не знайдено", foreground="#666")
        self.count_label.grid(
            row=1, column=0, columnspan=2, sticky="w", pady=(8, 0))

    # ------------------------------------------------------------ Дії GUI --
    def select_file(self):
        file_path = filedialog.askopenfilename(
            title="Виберіть файл протоколу",
            filetypes=[("Word документи", "*.docx"), ("Всі файли", "*.*")]
        )
        if file_path:
            self.input_file = file_path
            self.file_label.config(
                text=f"✓ {os.path.basename(file_path)}", foreground="#2e7d32")
            self.status_label.config(text="")

    def open_executors_manager(self):
        ExecutorsManager(self.root, self.executors, self._on_executors_saved)

    def _on_executors_saved(self, new_executors):
        self.executors = new_executors
        messagebox.showinfo("Готово", "Список виконавців оновлено.")

    def _show_full_text(self, event):
        sel = self.tree.selection()
        if not sel:
            return
        values = self.tree.item(sel[0], "values")
        messagebox.showinfo(f"Доручення № {values[0]}",
                            f"Виконавець: {values[2]}\nТермін: {values[1]}\n\n{values[3]}")

    # -------------------------------------------------------- Обробка тексту --
    def _parse_deadline(self, line):
        """Повертає термін виконання. Якщо в тексті конкретна дата - повертає
        її. Якщо термін вказано словами (напр. «червень-серпень 2026»,
        «серпень/вересень») - повертає цей текст так, як він написаний.
        Якщо «постійно» - повертає «постійно»."""
        date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', line)
        if date_match:
            return date_match.group(1)
        if 'постійно' in line.lower():
            return 'постійно'
        after_colon = re.split(
            r'ермін[^:]*:\s*', line, maxsplit=1, flags=re.IGNORECASE)
        if len(after_colon) > 1:
            text = after_colon[1].strip().rstrip('.').strip()
            if text:
                return text
        return 'Не вказано'

    def extract_protocol_data(self, doc_path):
        """Витягує список доручень з протоколу."""
        doc = docx.Document(doc_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', full_text)
        protocol_date = date_match.group(1) if date_match else "?"

        # Починаємо парсинг від слова "ВИРІШЕНО"
        decisions_start = -1
        for keyword in ['ВИРІШЕНО', 'вирішено', 'Вирішено']:
            pos = full_text.find(keyword)
            if pos != -1:
                decisions_start = pos
                break
        work_text = full_text[decisions_start:] if decisions_start != - \
            1 else full_text

        # Відсікаємо все, що починається з "Протокол вела" / "Протокол вів"
        end_match = re.search(r'Протокол\s+(вела|вів)',
                              work_text, re.IGNORECASE)
        if end_match:
            work_text = work_text[:end_match.start()]

        executor_patterns = build_executor_patterns(self.executors)

        decisions = []
        lines = work_text.split('\n')

        current_text = []
        current_executor = None
        current_deadline = None
        started = False

        def flush():
            if current_text:
                text_combined = re.sub(
                    r'\s+', ' ', ' '.join(current_text)).strip()
                if text_combined:
                    decisions.append({
                        'text': text_combined,
                        'executor': current_executor or 'Не вказано',
                        'date_given': protocol_date,
                        'deadline': current_deadline or 'Не вказано',
                    })

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            is_new_decision = False
            found_executor = None
            matched_pattern = None

            for pattern, canonical in executor_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    is_new_decision = True
                    found_executor = canonical
                    matched_pattern = pattern
                    break

            numbered_match = re.match(r'^(\d+)\.?\s+', line_stripped)
            if numbered_match and len(numbered_match.group(1)) <= 3:
                is_new_decision = True

            if is_new_decision:
                flush()

                clean_line = line_stripped
                if numbered_match:
                    clean_line = clean_line[numbered_match.end():].strip()
                if matched_pattern:
                    clean_line = re.sub(matched_pattern, '',
                                        clean_line, flags=re.IGNORECASE).strip()

                if clean_line:
                    clean_line = clean_line[0].upper() + clean_line[1:]

                current_text = [clean_line] if clean_line else []
                current_executor = found_executor
                current_deadline = None
                started = True

            elif started:
                if 'ермін' in line:
                    current_deadline = self._parse_deadline(line)
                else:
                    current_text.append(line_stripped)

        flush()
        decisions = self._split_meeting_items(decisions)
        return decisions

    def _split_meeting_items(self, decisions):
        """Підстраховка: якщо пункт про наступне засідання злипся з
        попереднім дорученням (через особливості форматування вихідного
        файлу), примусово розрізає текст на два окремих пункти."""
        meeting_re = re.compile(
            r'(Призначити (?:наступне|чергове) засідання.*)', re.IGNORECASE | re.DOTALL)

        result = []
        for item in decisions:
            match = meeting_re.search(item['text'])
            if match and match.start() > 0:
                before_text = item['text'][:match.start()].strip(' .')
                after_text = item['text'][match.start():].strip()

                if before_text:
                    before_item = dict(item)
                    before_item['text'] = before_text
                    result.append(before_item)

                if after_text:
                    after_text = after_text[0].upper() + after_text[1:]
                    after_item = dict(item)
                    after_item['text'] = after_text
                    after_item['executor'] = 'Всі присутні'
                    result.append(after_item)
            else:
                result.append(item)
        return result

    def sort_data(self, data):
        """Сортує доручення: спочатку за терміном (реальні дати за
        зростанням, потім нечислові терміни, потім «постійно»), в межах
        одного терміну - за виконавцем; пункт про наступне засідання
        завжди йде останнім у межах свого терміну. Після сортування
        нумерує доручення заново від 1."""

        def sort_key(item):
            deadline = item['deadline']
            try:
                if deadline not in ('Не вказано', 'постійно'):
                    date_val = datetime.strptime(deadline, '%d.%m.%Y')
                    category = 0
                    value = date_val.strftime('%Y%m%d')
                elif deadline == 'постійно':
                    category = 2
                    value = ''
                else:
                    category = 1
                    value = deadline
            except ValueError:
                category = 1
                value = deadline

            text_lower = item['text'].lower()
            is_meeting = (
                'наступне засідання' in text_lower or 'призначити наступне' in text_lower)
            meeting_priority = 1 if is_meeting else 0

            executor = item['executor'] if item['executor'] != 'Не вказано' else 'ЯЯЯ'

            return (category, value, meeting_priority, executor)

        sorted_list = sorted(data, key=sort_key)
        for i, item in enumerate(sorted_list, start=1):
            item['num'] = str(i)
        return sorted_list

    # --------------------------------------------------------------- Кнопки --
    def process_file(self):
        if not self.input_file:
            messagebox.showwarning(
                "Увага", "Спочатку виберіть файл протоколу!")
            return

        try:
            raw_data = self.extract_protocol_data(self.input_file)
            self.output_data = self.sort_data(raw_data)
            self.update_display()
            self.status_label.config(
                text=f"✓ Оброблено успішно: {len(self.output_data)} доручень", foreground="#2e7d32")
        except Exception as e:
            messagebox.showerror(
                "Помилка", f"Не вдалося обробити файл:\n{str(e)}")
            self.status_label.config(
                text="✗ Помилка обробки", foreground="#c62828")

    def update_display(self):
        self.tree.delete(*self.tree.get_children())
        for item in self.output_data:
            preview = item['text'] if len(
                item['text']) <= 140 else item['text'][:137] + "..."
            self.tree.insert("", "end", values=(
                item['num'], item['deadline'], item['executor'], preview))
        self.count_label.config(text=f"Знайдено доручень: {len(self.output_data)}  "
                                f"(подвійний клік - переглянути повний текст)")

    def save_result(self):
        if not self.output_data:
            messagebox.showwarning("Увага", "Спочатку обробіть файл!")
            return

        save_path = filedialog.asksaveasfilename(
            title="Зберегти результат",
            defaultextension=".docx",
            filetypes=[("Word документи", "*.docx"), ("Всі файли", "*.*")]
        )
        if save_path:
            try:
                self.create_table_document(save_path)
                messagebox.showinfo(
                    "Успіх", f"Таблицю збережено:\n{save_path}")
            except Exception as e:
                messagebox.showerror(
                    "Помилка", f"Не вдалося зберегти файл:\n{str(e)}")

    def clear_all(self):
        self.input_file = None
        self.output_data = []
        self.file_label.config(text="Файл не вибрано", foreground="#666")
        self.status_label.config(text="")
        self.tree.delete(*self.tree.get_children())
        self.count_label.config(text="Доручень ще не знайдено")

    # --------------------------------------------------------- Формування docx --
    def create_table_document(self, output_path):
        doc = docx.Document()

        section = doc.sections[0]
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Дата наступного засідання - перша реальна дата у відсортованому списку
        next_meeting_date = "Не вказано"
        for item in self.output_data:
            if re.match(r'^\d{2}\.\d{2}\.\d{4}$', item['deadline']):
                next_meeting_date = item['deadline']
                break

        for text, bold in [
            ("ІНФОРМАЦІЯ", True),
            ("з виконання протокольних доручень голови районної адміністрації Ільченко С.В.", False),
            (f"станом на {next_meeting_date}", False),
        ]:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = bold
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        col_widths = [
            Inches(0.3937),  # № з/п - 1 см
            Inches(5.2),     # Доручення
            Inches(1.3),     # Виконавці
            Inches(1.0),     # Термін надання
            Inches(1.0),     # Термін виконання
            Inches(1.9685),  # Примітки - 5 см
        ]
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = width

        headers = ['№ з/п', 'Доручення', 'Виконавці', 'Термін надання доручення',
                   'Термін виконання доручення', 'Примітки']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
            header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for item in self.output_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['num']
            self._format_cell(row_cells[0], WD_ALIGN_PARAGRAPH.CENTER)

            row_cells[1].text = item['text']
            self._format_cell(row_cells[1], WD_ALIGN_PARAGRAPH.LEFT)

            row_cells[2].text = item['executor']
            self._format_cell(row_cells[2], WD_ALIGN_PARAGRAPH.CENTER)

            row_cells[3].text = item['date_given']
            self._format_cell(row_cells[3], WD_ALIGN_PARAGRAPH.CENTER)

            row_cells[4].text = item['deadline']
            self._format_cell(row_cells[4], WD_ALIGN_PARAGRAPH.CENTER)

            row_cells[5].text = ''
            self._format_cell(row_cells[5], WD_ALIGN_PARAGRAPH.CENTER)

        doc.save(output_path)

    def _format_cell(self, cell, alignment):
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    root = tk.Tk()
    ui_font = get_ui_font()
    root.option_add("*Font", (ui_font, 10))
    app = ProtocolProcessor(root, ui_font)
    root.mainloop()


if __name__ == "__main__":
    main()
